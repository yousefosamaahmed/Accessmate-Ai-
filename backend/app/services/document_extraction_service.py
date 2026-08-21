# app/services/document_extraction_service.py

import csv
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


class DocumentExtractionService:
    """
    Extract readable text from supported document types.

    Supported:
    - TXT
    - CSV
    - DOCX
    - PDF

    Legacy .doc files are intentionally not handled here
    because python-docx does not support the old binary DOC format.
    """

    SUPPORTED_FILE_TYPES = {
        "txt",
        "csv",
        "docx",
        "pdf",
    }


    # ========================================================
    # MAIN ENTRY POINT
    # ========================================================

    def extract_text(
        self,
        file_path: str,
        file_type: str,
    ) -> str:
        normalized_file_type = (
            str(file_type or "")
            .lower()
            .strip()
            .replace(".", "")
        )

        if not normalized_file_type:
            raise ValueError(
                "Document file type is required"
            )

        if normalized_file_type == "txt":
            return self.extract_text_from_txt(
                file_path
            )

        if normalized_file_type == "csv":
            return self.extract_text_from_csv(
                file_path
            )

        if normalized_file_type == "docx":
            return self.extract_text_from_docx(
                file_path
            )

        if normalized_file_type == "pdf":
            return self.extract_text_from_pdf(
                file_path
            )

        if normalized_file_type == "doc":
            raise ValueError(
                "Legacy .doc files are not supported for text extraction. "
                "Please convert the file to .docx before uploading."
            )

        raise ValueError(
            f"Text extraction is not supported for file type: "
            f"{file_type}"
        )


    # ========================================================
    # FILE VALIDATION
    # ========================================================

    def _get_existing_file(
        self,
        file_path: str,
    ) -> Path:
        path = Path(
            file_path
        )

        if not path.exists():
            raise ValueError(
                "File not found"
            )

        if not path.is_file():
            raise ValueError(
                "Invalid file path"
            )

        return path


    # ========================================================
    # TXT
    # ========================================================

    def extract_text_from_txt(
        self,
        file_path: str,
    ) -> str:
        path = self._get_existing_file(
            file_path
        )

        # UTF-8 with BOM support first.
        try:
            text = path.read_text(
                encoding="utf-8-sig"
            )

        except UnicodeDecodeError:
            # Conservative Windows fallback.
            text = path.read_text(
                encoding="cp1252",
                errors="ignore",
            )

        return text.strip()


    # ========================================================
    # CSV
    # ========================================================

    def extract_text_from_csv(
        self,
        file_path: str,
    ) -> str:
        """
        Convert CSV rows into readable text while preserving
        column relationships.

        Example:

        Name | Age | Country
        Ahmed | 24 | Egypt

        This is much more useful for RAG than returning a raw
        comma-separated byte stream.
        """

        path = self._get_existing_file(
            file_path
        )

        rows: list[list[str]] = []


        # ----------------------------------------------------
        # Try common encodings
        # ----------------------------------------------------

        encodings = [
            "utf-8-sig",
            "utf-8",
            "cp1256",
            "cp1252",
        ]


        last_error: Exception | None = None


        for encoding in encodings:
            try:
                with path.open(
                    "r",
                    encoding=encoding,
                    newline="",
                ) as csv_file:

                    sample = csv_file.read(
                        8192
                    )

                    csv_file.seek(0)


                    # ----------------------------------------
                    # Detect delimiter when possible.
                    # ----------------------------------------

                    try:
                        dialect = (
                            csv.Sniffer()
                            .sniff(
                                sample,
                                delimiters=",;\t|",
                            )
                        )

                    except csv.Error:
                        dialect = (
                            csv.excel
                        )


                    reader = csv.reader(
                        csv_file,
                        dialect,
                    )


                    rows = [
                        [
                            str(cell).strip()
                            for cell in row
                        ]
                        for row in reader
                        if any(
                            str(cell).strip()
                            for cell in row
                        )
                    ]


                last_error = None
                break

            except UnicodeDecodeError as error:
                last_error = error
                continue


        if last_error is not None:
            raise ValueError(
                "Unable to decode CSV file"
            ) from last_error


        if not rows:
            return ""


        # ----------------------------------------------------
        # Build RAG-friendly text.
        # ----------------------------------------------------

        formatted_rows: list[str] = []


        for row in rows:
            formatted_rows.append(
                " | ".join(
                    row
                )
            )


        return "\n".join(
            formatted_rows
        ).strip()


    # ========================================================
    # DOCX
    # ========================================================

    def extract_text_from_docx(
        self,
        file_path: str,
    ) -> str:
        path = self._get_existing_file(
            file_path
        )


        try:
            document = DocxDocument(
                str(path)
            )

        except Exception as error:
            raise ValueError(
                "Unable to read DOCX document"
            ) from error


        text_parts: list[str] = []


        # ----------------------------------------------------
        # Paragraphs
        # ----------------------------------------------------

        for paragraph in document.paragraphs:
            paragraph_text = (
                paragraph.text.strip()
            )

            if paragraph_text:
                text_parts.append(
                    paragraph_text
                )


        # ----------------------------------------------------
        # Tables
        #
        # Important because many reports/CVs/business docs
        # keep meaningful information inside Word tables.
        # ----------------------------------------------------

        for table in document.tables:

            for row in table.rows:

                cells = [
                    cell.text.strip()
                    for cell in row.cells
                ]


                cells = [
                    cell
                    for cell in cells
                    if cell
                ]


                if cells:
                    text_parts.append(
                        " | ".join(
                            cells
                        )
                    )


        return "\n".join(
            text_parts
        ).strip()


    # ========================================================
    # PDF
    # ========================================================

    def extract_text_from_pdf(
        self,
        file_path: str,
    ) -> str:
        path = self._get_existing_file(
            file_path
        )


        try:
            reader = PdfReader(
                str(path)
            )

        except Exception as error:
            raise ValueError(
                "Unable to read PDF document"
            ) from error


        pages_text: list[str] = []


        for page_number, page in enumerate(
            reader.pages,
            start=1,
        ):
            try:
                page_text = (
                    page.extract_text()
                )

            except Exception:
                page_text = None


            if not page_text:
                continue


            cleaned_page_text = (
                page_text.strip()
            )


            if not cleaned_page_text:
                continue


            # Page marker is useful for RAG context.
            pages_text.append(
                f"[Page {page_number}]\n"
                f"{cleaned_page_text}"
            )


        return "\n\n".join(
            pages_text
        ).strip()


    # ========================================================
    # SUPPORT CHECK
    # ========================================================

    def is_supported(
        self,
        file_type: str,
    ) -> bool:
        normalized_file_type = (
            str(file_type or "")
            .lower()
            .strip()
            .replace(".", "")
        )

        return (
            normalized_file_type
            in self.SUPPORTED_FILE_TYPES
        )